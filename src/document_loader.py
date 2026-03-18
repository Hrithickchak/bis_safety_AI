"""
Document Loader – ingest safety procedure documents.

Supports: plain-text (.txt), PDF (.pdf), Word (.docx), and HTML (.html/.htm).
Extracts raw text and attempts to split into logical sections.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from src.models import ProcedureDocument, ProcedureSection


# ---------------------------------------------------------------------------
# Low-level readers
# ---------------------------------------------------------------------------

def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("Install pypdf: pip install pypdf")
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        raise ImportError("Install python-docx: pip install python-docx")
    doc = docx.Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _read_html(path: Path) -> str:
    try:
        import html2text
    except ImportError:
        raise ImportError("Install html2text: pip install html2text")
    raw = path.read_text(encoding="utf-8", errors="replace")
    converter = html2text.HTML2Text()
    converter.ignore_links = False
    converter.body_width = 0
    return converter.handle(raw)


# Mapping of file extensions to reader functions
_READERS = {
    ".txt": _read_txt,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".html": _read_html,
    ".htm": _read_html,
}


# ---------------------------------------------------------------------------
# Section splitting
# ---------------------------------------------------------------------------

# Patterns that look like section headings (numbered or ALL-CAPS lines)
_HEADING_PATTERNS = [
    # Numbered sections: "1. Purpose", "2.3 Scope", "Section 4 – Training"
    re.compile(
        r"^(?:Section\s+)?\d+(?:\.\d+)*[\.\)\s\-–—]+\s*[A-Z].*",
        re.MULTILINE,
    ),
    # ALL-CAPS headings (at least 3 words, stand-alone line)
    re.compile(r"^[A-Z][A-Z\s\-/&]{8,}$", re.MULTILINE),
    # Markdown-style headings
    re.compile(r"^#{1,4}\s+.+", re.MULTILINE),
]


def _split_into_sections(text: str) -> list[ProcedureSection]:
    """Split document text into sections based on heading patterns."""
    # Collect heading positions
    headings: list[tuple[int, int, str]] = []
    for pat in _HEADING_PATTERNS:
        for m in pat.finditer(text):
            headings.append((m.start(), m.end(), m.group().strip()))

    if not headings:
        # No headings detected → return the whole text as a single section
        return [ProcedureSection(heading="(Full Document)", text=text.strip())]

    # Sort by position
    headings.sort(key=lambda x: x[0])

    # Deduplicate overlapping headings (keep the first)
    deduped: list[tuple[int, int, str]] = []
    last_end = -1
    for start, end, heading in headings:
        if start >= last_end:
            deduped.append((start, end, heading))
            last_end = end
    headings = deduped

    sections: list[ProcedureSection] = []
    for i, (start, end, heading) in enumerate(headings):
        next_start = headings[i + 1][0] if i + 1 < len(headings) else len(text)
        body = text[end:next_start].strip()
        if body or heading:
            sections.append(ProcedureSection(heading=heading, text=body))

    # Capture any text before the first heading
    preamble = text[: headings[0][0]].strip()
    if preamble:
        sections.insert(
            0, ProcedureSection(heading="(Preamble)", text=preamble)
        )

    return sections


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def load_procedure(
    path: str | Path,
    title: Optional[str] = None,
) -> ProcedureDocument:
    """
    Load a safety procedure document from disk.

    Parameters
    ----------
    path : str or Path
        Path to the procedure file (.txt, .pdf, .docx, .html).
    title : str, optional
        Override title; defaults to the filename stem.

    Returns
    -------
    ProcedureDocument
        Parsed document with raw text and sections.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Procedure file not found: {path}")

    suffix = path.suffix.lower()
    reader = _READERS.get(suffix)
    if reader is None:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            f"Supported: {', '.join(_READERS.keys())}"
        )

    raw_text = reader(path)
    sections = _split_into_sections(raw_text)

    return ProcedureDocument(
        title=title or path.stem.replace("_", " ").title(),
        source_path=str(path),
        raw_text=raw_text,
        sections=sections,
    )


def load_all_procedures(directory: str | Path) -> list[ProcedureDocument]:
    """Load all procedure files from a directory."""
    directory = Path(directory)
    docs: list[ProcedureDocument] = []
    for ext in _READERS:
        for file_path in sorted(directory.glob(f"*{ext}")):
            docs.append(load_procedure(file_path))
    return docs
