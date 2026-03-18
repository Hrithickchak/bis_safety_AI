"""
Regulatory Store – vector database for regulatory texts with RAG retrieval.

Loads regulation documents, splits them into chunks, embeds them using
Google Gemini embeddings, and stores them in a FAISS vector store.
Provides semantic search to retrieve the most relevant regulatory clauses
for a given query or procedure excerpt.
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import time
from pathlib import Path
from typing import Optional

from config import get_settings

logger = logging.getLogger(__name__)

# Free-tier Gemini embedding limit: 100 requests/min.
# We batch documents and pause between batches to stay under.
_EMBED_BATCH_SIZE = 20
_EMBED_BATCH_DELAY = 4.0  # seconds between batches


# ---------------------------------------------------------------------------
# Text splitting
# ---------------------------------------------------------------------------

def _split_text(
    text: str,
    chunk_size: int = 800,
    chunk_overlap: int = 200,
    source: str = "",
) -> list[dict]:
    """
    Split text into overlapping chunks with metadata.

    Returns a list of dicts with keys: 'text', 'metadata'.
    """
    chunks: list[dict] = []
    lines = text.split("\n")
    current_chunk: list[str] = []
    current_len = 0
    chunk_idx = 0

    for line_num, line in enumerate(lines, 1):
        line_len = len(line) + 1  # +1 for newline
        if current_len + line_len > chunk_size and current_chunk:
            chunk_text = "\n".join(current_chunk)
            chunks.append({
                "text": chunk_text,
                "metadata": {
                    "source": source,
                    "chunk_index": chunk_idx,
                    "start_line": line_num - len(current_chunk),
                },
            })
            chunk_idx += 1
            # Keep overlap
            overlap_chars = 0
            overlap_lines: list[str] = []
            for prev_line in reversed(current_chunk):
                overlap_chars += len(prev_line) + 1
                overlap_lines.insert(0, prev_line)
                if overlap_chars >= chunk_overlap:
                    break
            current_chunk = overlap_lines
            current_len = sum(len(l) + 1 for l in current_chunk)

        current_chunk.append(line)
        current_len += line_len

    # Last chunk
    if current_chunk:
        chunk_text = "\n".join(current_chunk)
        chunks.append({
            "text": chunk_text,
            "metadata": {
                "source": source,
                "chunk_index": chunk_idx,
                "start_line": len(lines) - len(current_chunk) + 1,
            },
        })

    return chunks


# ---------------------------------------------------------------------------
# Regulatory Vector Store
# ---------------------------------------------------------------------------

class RegulatoryStore:
    """
    Vector store backed by FAISS for regulatory text retrieval.

    Uses Google Gemini embeddings for vectorisation.

    Provides methods to:
    - Ingest regulation text files into the vector store.
    - Perform semantic search to retrieve relevant regulatory clauses.
    - Persist and reload the index from disk.
    """

    def __init__(
        self,
        persist_dir: Optional[str | Path] = None,
        force_rebuild: bool = False,
    ):
        settings = get_settings()
        self.persist_dir = Path(
            persist_dir or settings.vectorstore_dir
        )
        self.persist_dir.mkdir(parents=True, exist_ok=True)
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.top_k = settings.retrieval_top_k

        self._embeddings = None
        self._vectorstore = None
        self._documents: list[dict] = []
        self._force_rebuild = force_rebuild

    # ---- lazy initialization ----

    def _get_embeddings(self):
        """Lazily initialize the Gemini embedding model."""
        if self._embeddings is None:
            from langchain_google_genai import GoogleGenerativeAIEmbeddings
            settings = get_settings()
            self._embeddings = GoogleGenerativeAIEmbeddings(
                model=settings.gemini_embedding_model,
                google_api_key=settings.google_api_key,
            )
        return self._embeddings

    def _get_or_create_store(self):
        """Load existing vectorstore from disk, or create a new one."""
        if self._vectorstore is not None:
            return self._vectorstore

        index_path = self.persist_dir / "index.faiss"

        if index_path.exists() and not self._force_rebuild:
            from langchain_community.vectorstores import FAISS
            self._vectorstore = FAISS.load_local(
                str(self.persist_dir),
                self._get_embeddings(),
                allow_dangerous_deserialization=True,
            )
        return self._vectorstore

    # ---- ingestion ----

    def ingest_file(self, file_path: str | Path, source_label: str = "") -> int:
        """
        Read a regulation file and add its chunks to the vector store.

        Supports .txt, .pdf, .docx, and .html/.htm files.

        Parameters
        ----------
        file_path : path
            Path to a regulation file.
        source_label : str
            Label for the source (e.g. 'OSHA 1910.178'). Defaults to
            the filename stem.

        Returns
        -------
        int
            Number of chunks added.
        """
        file_path = Path(file_path)
        label = source_label or file_path.stem.replace("_", " ").upper()
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            text = self._read_pdf(file_path)
        elif suffix == ".docx":
            text = self._read_docx(file_path)
        elif suffix in (".html", ".htm"):
            text = self._read_html(file_path)
        else:
            text = file_path.read_text(encoding="utf-8", errors="replace")

        return self.ingest_text(text, source_label=label)

    @staticmethod
    def _read_pdf(path: Path) -> str:
        """Extract text from a PDF file."""
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)

    @staticmethod
    def _read_docx(path: Path) -> str:
        """Extract text from a DOCX file."""
        import docx
        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)

    @staticmethod
    def _read_html(path: Path) -> str:
        """Extract text from an HTML file."""
        import html2text
        raw = path.read_text(encoding="utf-8", errors="replace")
        converter = html2text.HTML2Text()
        converter.ignore_links = False
        converter.body_width = 0
        return converter.handle(raw)

    def ingest_text(self, text: str, source_label: str = "") -> int:
        """
        Ingest raw regulation text into the vector store.

        Parameters
        ----------
        text : str
            Raw regulation text.
        source_label : str
            Source identifier.

        Returns
        -------
        int
            Number of chunks added.
        """
        chunks = _split_text(
            text,
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            source=source_label,
        )
        if not chunks:
            return 0

        self._documents.extend(chunks)

        from langchain_community.vectorstores import FAISS
        from langchain_core.documents import Document

        lc_docs = [
            Document(page_content=c["text"], metadata=c["metadata"])
            for c in chunks
        ]

        # Process in small batches to respect Gemini free-tier rate limits
        for batch_start in range(0, len(lc_docs), _EMBED_BATCH_SIZE):
            batch = lc_docs[batch_start : batch_start + _EMBED_BATCH_SIZE]
            batch_num = batch_start // _EMBED_BATCH_SIZE + 1
            total_batches = (len(lc_docs) + _EMBED_BATCH_SIZE - 1) // _EMBED_BATCH_SIZE

            # Retry loop for rate-limit errors
            for attempt in range(5):
                try:
                    if self._vectorstore is None:
                        self._vectorstore = FAISS.from_documents(
                            batch, self._get_embeddings()
                        )
                    else:
                        self._vectorstore.add_documents(batch)
                    break  # success
                except Exception as e:
                    if "429" in str(e) or "RESOURCE_EXHAUSTED" in str(e):
                        wait = _EMBED_BATCH_DELAY * (attempt + 1)
                        logger.warning(
                            "Rate limited (batch %d/%d), waiting %.0fs...",
                            batch_num, total_batches, wait,
                        )
                        time.sleep(wait)
                    else:
                        raise

            # Pause between batches to avoid hitting the rate limit
            if batch_start + _EMBED_BATCH_SIZE < len(lc_docs):
                logger.info(
                    "Embedded batch %d/%d (%d docs), pausing %.0fs...",
                    batch_num, total_batches, len(batch), _EMBED_BATCH_DELAY,
                )
                time.sleep(_EMBED_BATCH_DELAY)

        return len(chunks)

    def ingest_directory(self, directory: str | Path) -> int:
        """Ingest all supported regulation files (.txt, .pdf, .docx, .html) from a directory."""
        directory = Path(directory)
        supported_extensions = {".txt", ".pdf", ".docx", ".html", ".htm"}
        total = 0
        for f in sorted(directory.iterdir()):
            if f.suffix.lower() in supported_extensions:
                n = self.ingest_file(f)
                total += n
        return total

    # ---- persistence ----

    def save(self) -> None:
        """Persist the vector store to disk."""
        if self._vectorstore is not None:
            self._vectorstore.save_local(str(self.persist_dir))
            # Also save document metadata
            meta_path = self.persist_dir / "documents_meta.json"
            with open(meta_path, "w", encoding="utf-8") as f:
                json.dump(self._documents, f, indent=2)

    def load(self) -> bool:
        """Load the vector store from disk. Returns True if successful."""
        store = self._get_or_create_store()
        return store is not None

    # ---- retrieval ----

    def search(
        self,
        query: str,
        top_k: Optional[int] = None,
        filter_source: Optional[str] = None,
    ) -> list[dict]:
        """
        Perform semantic search against the regulatory vector store.

        Parameters
        ----------
        query : str
            Search query (e.g. a procedure excerpt or topic).
        top_k : int, optional
            Number of results to return.
        filter_source : str, optional
            Filter results by source label.

        Returns
        -------
        list[dict]
            List of dicts with 'text', 'metadata', and 'score'.
        """
        store = self._get_or_create_store()
        if store is None:
            return []

        k = top_k or self.top_k

        results = store.similarity_search_with_score(query, k=k)

        output: list[dict] = []
        for doc, score in results:
            if filter_source and doc.metadata.get("source", "") != filter_source:
                continue
            output.append({
                "text": doc.page_content,
                "metadata": doc.metadata,
                "score": float(score),
            })

        return output

    def retrieve_for_procedure(
        self,
        procedure_text: str,
        standards: list[str] | None = None,
        top_k: int | None = None,
    ) -> list[dict]:
        """
        Retrieve regulatory clauses relevant to a procedure.

        Breaks the procedure into key sentences and retrieves
        matching regulation chunks.

        Parameters
        ----------
        procedure_text : str
            Full text of the safety procedure.
        standards : list[str], optional
            If provided, only retrieve from these standard sources.
        top_k : int, optional
            Override the default number of results.

        Returns
        -------
        list[dict]
            Deduplicated list of relevant regulatory chunks.
        """
        k = top_k or self.top_k

        # Extract key sentences (those containing action verbs / requirements)
        import re
        sentences = re.split(r"[.!?\n]+", procedure_text)
        key_sentences = [
            s.strip() for s in sentences
            if len(s.strip()) > 30  # skip very short fragments
        ]

        # Limit to avoid too many queries
        if len(key_sentences) > 20:
            # Sample evenly from the document
            step = len(key_sentences) // 20
            key_sentences = key_sentences[::step][:20]

        seen_texts: set[str] = set()
        all_results: list[dict] = []

        for sentence in key_sentences:
            results = self.search(sentence, top_k=k // 2 or 4)
            for r in results:
                text_hash = hashlib.md5(
                    r["text"].encode()
                ).hexdigest()
                if text_hash not in seen_texts:
                    seen_texts.add(text_hash)
                    all_results.append(r)

        # Also do a broad search with the full text (truncated)
        broad_results = self.search(
            procedure_text[:2000], top_k=k
        )
        for r in broad_results:
            text_hash = hashlib.md5(r["text"].encode()).hexdigest()
            if text_hash not in seen_texts:
                seen_texts.add(text_hash)
                all_results.append(r)

        # Sort by relevance score (lower = more similar for FAISS L2)
        all_results.sort(key=lambda x: x["score"])

        # Filter by standard if specified
        if standards:
            standards_lower = [s.lower() for s in standards]
            all_results = [
                r for r in all_results
                if any(
                    std in r["metadata"].get("source", "").lower()
                    for std in standards_lower
                )
            ] or all_results  # fallback to all if filter removes everything

        return all_results[:top_k or self.top_k * 4]

    @property
    def document_count(self) -> int:
        """Number of documents ingested."""
        return len(self._documents)
